from docx import Document

def main():
    document = Document('sample.docx')
    document.add_picture('yukidaruma.png')
    count = 0
    for data in document.paragraphs:
        count += len(data.text)

    print(f'sample.docxの中の文字数は{count}個です。')

    document.save('sample_after.docx')

if __name__ == '__main__':
    main()